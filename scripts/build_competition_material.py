from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "开题报告附加材料release.docx"
ASSET_DIR = ROOT / "docs" / "assets" / "readme"
GENERATED_DIR = ROOT / ".codex-docx-assets"

# narrative_proposal preset + named CJK readability overrides.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "0F172A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GREEN = "16835F"
GREEN_DARK = "166534"
GREEN_LIGHT = "ECFDF5"
GOLD = "D97706"
GOLD_LIGHT = "FFF8E8"
PURPLE = "6D28D9"
INK = "1F2937"
MUTED = "64748B"
LINE = "D9E2EC"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
CJK_BODY_FONT = "Microsoft YaHei"
CJK_HEADING_FONT = "宋体"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, color: str = INK, bold: bool | None = None,
                 italic: bool | None = None, ascii_font: str = "Calibri",
                 east_asia_font: str = CJK_BODY_FONT) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    assert sum(widths) == CONTENT_DXA, (widths, sum(widths))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_HEADING_FONT)
        style._element.rPr.rFonts.attrib.pop(qn("w:eastAsiaTheme"), None)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    # Make every Chinese title style explicit so Word cannot resolve the
    # major-East-Asia theme to MS Gothic on another machine.
    for name in ["Title", "Subtitle", *[f"Heading {i}" for i in range(1, 10)]]:
        style = styles[name]
        r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), CJK_HEADING_FONT)
        r_fonts.attrib.pop(qn("w:eastAsiaTheme"), None)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_BODY_FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("青少年 C++ 算法可视化教学系统")
    set_run_font(run, size=9, color=MUTED, bold=True, east_asia_font=CJK_HEADING_FONT)
    run = hp.add_run("  |  开题报告附加材料 · 参赛项目佐证")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    def next_id(tag: str, attr: str) -> int:
        values = [int(x.get(qn(attr))) for x in numbering.findall(qn(tag)) if x.get(qn(attr))]
        return (max(values) + 1) if values else 1

    def make(fmt: str, text: str) -> int:
        abstract_id = next_id("w:abstractNum", "w:abstractNumId")
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num_id = next_id("w:num", "w:numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)
        return num_id

    return make("bullet", "•"), make("decimal", "%1.")


def add_list_item(doc: Document, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, size=11, color=INK, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_body(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        set_run_font(run, size=11, color=INK, bold=True)
        run = p.add_run(text[len(bold_lead):])
        set_run_font(run, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_callout(doc: Document, label: str, text: str, fill: str = GREEN_LIGHT, color: str = GREEN_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(label + "  ")
    set_run_font(r, size=11, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_picture(doc: Document, path: Path, width: float, caption: str, alt: str, figure_no: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    for doc_pr in run._element.xpath(".//wp:docPr"):
        doc_pr.set("descr", alt)
        doc_pr.set("title", alt)
    cp = doc.add_paragraph(style="Caption")
    r = cp.add_run(f"图 {figure_no}  {caption}")
    set_run_font(r, size=9, color=MUTED)


def add_metric_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_repeat_table_header(table.rows[0])
    metrics = [("9", "算法章节"), ("84", "课程课时"), ("201", "唯一题目"), ("35", "动画课时")]
    for i, (value, label) in enumerate(metrics):
        top = table.cell(0, i)
        bottom = table.cell(1, i)
        set_cell_shading(top, LIGHT)
        set_cell_shading(bottom, NAVY if i == 0 else GREEN if i == 1 else BLUE if i == 2 else GOLD)
        for paragraph in top.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
        r = top.paragraphs[0].add_run(label)
        set_run_font(r, size=10, color=INK, bold=True)
        for paragraph in bottom.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
        r = bottom.paragraphs[0].add_run(value)
        set_run_font(r, size=24, color=WHITE, bold=True)


def add_comparison_table(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1650, 3555, 4155])
    set_repeat_table_header(table.rows[0])
    headers = ["评委关注", "常见单点方案", "本项目的回答"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=WHITE, bold=True)
    rows = [
        ("是否真能学会", "只看动画或只刷题，理解与编码脱节", "动画步骤、C++ 变量、练习测试点使用一致知识映射"),
        ("是否可规模化", "每个知识点重复手工制作", "统一 Manim 主题、组件、脚本和课时数据约定"),
        ("是否能迭代", "课程发布后缺少反馈证据", "提交、通过率、薄弱题与错误类型回流教师看板"),
    ]
    for row_i, row in enumerate(rows, 1):
        for col_i, text in enumerate(row):
            cell = table.cell(row_i, col_i)
            if col_i == 0:
                set_cell_shading(cell, GREEN_LIGHT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=GREEN_DARK if col_i == 0 else INK, bold=(col_i == 0))


def add_curriculum_table(doc: Document) -> None:
    data = [
        ("1", "高精度计算", "ready", "11", "33", "11"),
        ("2", "数据排序", "ready", "6", "18", "6"),
        ("3", "递推算法", "building", "9", "27", "9"),
        ("4", "递归算法", "ready", "9", "27", "9"),
        ("5", "搜索与回溯", "building", "10", "20", "0"),
        ("6", "贪心算法", "ready", "10", "20", "0"),
        ("7", "分治算法", "ready", "8", "16", "0"),
        ("8", "广度优先搜索", "ready", "9", "18", "0"),
        ("9", "动态规划", "ready", "12", "24", "0"),
        ("合计", "", "7 ready / 2 building", "84", "203", "35"),
    ]
    table = doc.add_table(rows=1 + len(data), cols=6)
    table.style = "Table Grid"
    set_table_geometry(table, [720, 2500, 1900, 1320, 1460, 1460])
    headers = ["章", "主题", "状态", "课时", "练习", "动画"]
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=WHITE, bold=True)
    for r_i, row in enumerate(data, 1):
        for c_i, text in enumerate(row):
            cell = table.cell(r_i, c_i)
            if r_i == len(data):
                set_cell_shading(cell, GREEN_LIGHT)
            elif r_i % 2 == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            color = GREEN_DARK if text == "ready" else GOLD if text == "building" else INK
            set_run_font(r, size=9, color=color, bold=(r_i == len(data) or c_i == 2))


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill: str):
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=fnt)
    x = left + (right - left - (bbox[2] - bbox[0])) / 2
    y = top + (bottom - top - (bbox[3] - bbox[1])) / 2 - 3
    draw.text((x, y), text, font=fnt, fill=fill)


def rounded_box(draw, xy, fill, outline, width=4, radius=22):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#16835f", width=8):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    draw.polygon([(ex, ey), (ex - 20, ey - 12), (ex - 20, ey + 12)], fill=fill)


def create_learning_loop_png(path: Path) -> None:
    img = Image.new("RGB", (1800, 620), "#f8fafc")
    d = ImageDraw.Draw(img)
    center_text(d, (0, 22, 1800, 78), "从“看懂”到“会写”，再回到“精准教”", font(37, True), "#0f172a")
    center_text(d, (0, 80, 1800, 120), "动画、代码、练习、反馈与教学诊断使用同一知识映射", font(22), "#64748b")
    items = [
        ("① 动画理解", "状态 / 指针 / 调用栈", "#0f172a", "#0f172a", "#ffffff"),
        ("② 代码映射", "变量与步骤同步", "#ffffff", "#16a34a", "#166534"),
        ("③ 在线练习", "C++17 / 自动保存", "#ffffff", "#d97706", "#9a6700"),
        ("④ 即时判题", "测试点 / 错误类型", "#ffffff", "#2563eb", "#1d4ed8"),
        ("⑤ 分层反馈", "提示 / 题解 / 错题", "#ffffff", "#7c3aed", "#6d28d9"),
        ("⑥ 教师诊断", "薄弱点 / 教学建议", "#16835f", "#16835f", "#ffffff"),
    ]
    # Four balanced vertical bands: heading, process, feedback, and note.
    x0, y0, w, h, gap = 55, 145, 245, 160, 45
    for i, (title, sub, fill, outline, text_color) in enumerate(items):
        x = x0 + i * (w + gap)
        rounded_box(d, (x, y0, x + w, y0 + h), fill, outline, width=5)
        center_text(d, (x, y0 + 20, x + w, y0 + 83), title, font(28, True), text_color)
        center_text(d, (x, y0 + 88, x + w, y0 + 136), sub, font(20), "#d1fae5" if fill != "#ffffff" else "#52606d")
        if i < len(items) - 1:
            arrow(d, (x + w + 5, y0 + h // 2), (x + w + gap - 5, y0 + h // 2))
    d.arc((160, 280, 1650, 490), start=5, end=175, fill="#16835f", width=8)
    d.polygon([(166, 394), (190, 379), (190, 409)], fill="#16835f")
    center_text(d, (0, 395, 1800, 445), "数据反馈驱动内容迭代与下一次教学", font(27, True), "#16835f")
    center_text(d, (0, 520, 1800, 560), "区别于只播放动画或只提供题库的单点工具，本项目交付的是完整学习闭环。", font(20), "#52606d")
    img.save(path, quality=95)


def create_architecture_png(path: Path) -> None:
    img = Image.new("RGB", (1800, 750), "#f8fafc")
    d = ImageDraw.Draw(img)
    rounded_box(d, (55, 35, 1745, 140), "#0f172a", "#0f172a", radius=26)
    center_text(d, (55, 48, 1745, 100), "青少年 C++ 算法可视化教学系统", font(38, True), "#ffffff")
    center_text(d, (55, 100, 1745, 132), "内容生产、学习交互、在线判题、数据反馈四层协同", font(22), "#d1fae5")
    columns = [
        ("学习交互层", "React + TypeScript", ["章节 / 课时 / 进度", "动画、讲义、代码", "Monaco C++ 编辑器", "分层提示与复盘"], "#0f766e"),
        ("课程内容层", "Manim + 数据模型", ["统一动画视觉组件", "9 章 84 课时", "35 个动画课时", "201 个唯一题目"], "#2563eb"),
        ("服务与判题层", "FastAPI + Adapter", ["提交与结果 API", "C++17 测试点", "Local / Docker", "状态与错误分类"], "#d97706"),
        ("数据反馈层", "SQLite + 聚合统计", ["提交记录", "通过率", "薄弱题排行", "教学建议"], "#7c3aed"),
    ]
    x0, y0, w, h, gap = 65, 210, 385, 370, 45
    for i, (title, subtitle, bullets, color) in enumerate(columns):
        x = x0 + i * (w + gap)
        rounded_box(d, (x, y0, x + w, y0 + h), "#ffffff", color, width=5)
        center_text(d, (x, y0 + 20, x + w, y0 + 75), title, font(30, True), color)
        center_text(d, (x, y0 + 75, x + w, y0 + 120), subtitle, font(22), "#334155")
        yy = y0 + 150
        for bullet in bullets:
            d.text((x + 45, yy), "• " + bullet, font=font(21), fill="#475569")
            yy += 48
        if i < len(columns) - 1:
            arrow(d, (x + w + 6, y0 + h // 2), (x + w + gap - 6, y0 + h // 2), fill="#64748b", width=7)
    rounded_box(d, (380, 635, 1420, 705), "#ecfdf5", "#16835f", width=3, radius=30)
    center_text(d, (380, 645, 1420, 695), "模块化接口使内容、判题隔离与平台接入可独立演进", font(24, True), "#166534")
    img.save(path, quality=95)


def add_preview_strip(doc: Document) -> None:
    paths = [
        ROOT / "media" / "previews" / "quick_sort_preview.png",
        ROOT / "media" / "previews" / "recursion_call_stack_preview.png",
        ROOT / "media" / "previews" / "recurrence_number_tower_preview.png",
    ]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    alts = ["快速排序分区动画", "递归调用栈动画", "数塔递推状态动画"]
    for path, alt in zip(paths, alts):
        r = p.add_run()
        r.add_picture(str(path), width=Inches(2.08))
        for doc_pr in r._element.xpath(".//wp:docPr"):
            doc_pr.set("descr", alt)
            doc_pr.set("title", alt)
    cp = doc.add_paragraph(style="Caption")
    r = cp.add_run("图 8  跨章节动画视觉语言：快速排序、递归调用栈与数塔递推")
    set_run_font(r, size=9, color=MUTED)


def build() -> None:
    GENERATED_DIR.mkdir(exist_ok=True)
    loop_png = GENERATED_DIR / "learning-loop.png"
    architecture_png = GENERATED_DIR / "system-architecture.png"
    create_learning_loop_png(loop_png)
    create_architecture_png(architecture_png)

    doc = Document()
    configure_document(doc)
    bullet_id, decimal_id = create_numbering(doc)

    # Page 1: proposal_centerpiece cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("比赛参赛项目 · 开题报告附加材料")
    set_run_font(r, size=12, color=GREEN, bold=True, east_asia_font=CJK_HEADING_FONT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("青少年 C++ 算法可视化教学系统")
    set_run_font(r, size=25, color=NAVY, bold=True, east_asia_font=CJK_HEADING_FONT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("让抽象算法看得见、写得出、练得会")
    set_run_font(r, size=15, color=DARK_BLUE, bold=True, east_asia_font=CJK_HEADING_FONT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("可运行全栈 MVP | 算法动画 × 在线判题 × 学习数据闭环")
    set_run_font(r, size=10.5, color=MUTED, east_asia_font=CJK_HEADING_FONT)

    add_picture(doc, ASSET_DIR / "lesson-page.png", 6.45,
                "课程页将算法动画、理解路径、C++ 代码与练习入口放在同一学习空间",
                "高精度加法课程动画页", 1)
    add_metric_table(doc)
    add_callout(doc, "一句话价值", "用可复用的算法动画把抽象过程讲清楚，用在线判题把“看懂”转化为“会写”，再用学习数据帮助教师把下一次课讲得更准。")

    # Page 2: executive summary.
    doc.add_page_break()
    doc.add_heading("一、评委导读：为什么值得进入下一轮", level=1)
    add_body(doc, "本项目面向青少年 C++ 与信息学竞赛入门场景，聚焦两个长期痛点：一是算法中的状态、指针、调用栈和边界过程难以仅靠板书讲清；二是高质量动画课程制作成本高、练习与教学反馈彼此割裂。项目已经实现 React 学习前台、Manim 动画生产、FastAPI 判题服务与 SQLite 数据看板，形成可本地运行、可展示、可继续扩展的完整 MVP。")
    add_callout(doc, "建议评委先看", "项目的核心创新不在某一张动画或某一道题，而在同一知识点从动画、代码、练习、判题到教师诊断的全链路一致映射。", fill=GOLD_LIGHT, color=GOLD)
    add_comparison_table(doc)
    doc.add_heading("项目落地佐证", level=2)
    for text in [
        "当前课程数据实际包含 9 个章节、84 个课时、203 个课时练习引用（去重后 201 个唯一题目）；35 个课时已接入 Manim 动画。",
        "7 个章节状态为 ready，2 个章节处于持续建设；高精度、排序、递推与递归已形成动画课时样板。",
        "在线 C++17 编辑、自动保存、提交、测试点结果、提交历史、进度联动与教师统计看板均可运行。",
        "Judge Adapter 已抽象 LocalJudge / DockerJudge，两类判题实现可替换，便于从本地演示过渡到隔离部署。",
    ]:
        add_list_item(doc, text, bullet_id)
    doc.add_heading("申报结论", level=2)
    add_body(doc, "项目已越过“只有创意或界面原型”的早期阶段：内容、交互、判题、数据和演示脚本均有代码与界面证据。参赛价值在于把教学创新与工程落地放在同一套系统中验证，并为规模化内容生产留下清晰接口。")

    # Page 3: innovation loop.
    doc.add_page_break()
    doc.add_heading("二、核心创新：把单点工具连成学习闭环", level=1)
    add_picture(doc, loop_png, 6.5,
                "动画理解—代码映射—在线练习—即时判题—分层反馈—教师诊断",
                "算法学习闭环图", 2)
    innovations = [
        ("创新 1｜动画不是终点，而是代码理解的入口。", "动画中的数组格、指针、调用栈、状态转移与递归树继续映射到 C++ 变量、循环边界与测试点，减少“动画看懂了，代码仍不会写”的断层。"),
        ("创新 2｜分层反馈替代直接给答案。", "练习先提示思路，再提示关键变量与边界，最后才进入核心实现；同时提供题解要点与常见错误，为青少年提供适度支架。"),
        ("创新 3｜教师看板把错误转化为教学动作。", "系统聚合通过率、薄弱题、错误类型与平均尝试次数，把学生提交转化为下一次课的复讲重点与练习梯度建议。"),
        ("创新 4｜内容生产流程可复用。", "统一 Manim 主题、数组组件、预览图、场景命名与逐镜头脚本，让新增知识点能够复用已有视觉语言和工程流程。"),
    ]
    for title, body in innovations:
        doc.add_heading(title, level=2)
        add_body(doc, body)
    add_callout(doc, "评委看点", "创新点均有对应运行证据：图 4 证明学生端闭环，图 5 证明在线判题，图 6 证明数据回流；并非仅停留在概念描述。")

    # Page 4: architecture.
    doc.add_page_break()
    doc.add_heading("三、系统架构与工程创新", level=1)
    add_picture(doc, architecture_png, 6.5,
                "学习交互层、课程内容层、服务判题层与数据反馈层协同",
                "系统四层架构图", 3)
    doc.add_heading("模块化设计带来的可扩展性", level=2)
    for text in [
        "前端：React + TypeScript + Vite 负责课程导航、进度、视频、讲义、练习与教师看板。",
        "内容：Manim 负责算法动画生产，课程数据模型统一承载课时、示例、练习与媒体路径。",
        "服务：FastAPI 提供提交、判题、历史、统计与能力查询接口。",
        "判题：JudgeService 将业务层与执行环境隔离，可在 LocalJudge、DockerJudge 或外部 OJ Adapter 之间切换。",
        "数据：SQLite 记录提交并生成聚合统计，浏览器本地状态在后端不可用时提供体验兜底。",
    ]:
        add_list_item(doc, text, bullet_id, bold_prefix=text.split("：")[0] + "：")
    doc.add_heading("方案深度佐证", level=2)
    add_body(doc, "项目不是把多个开源组件简单拼接：课程内容模型决定动画与练习的映射，Judge Adapter 约束判题边界，统计看板又从提交记录计算教学指标。每一层都围绕同一教学闭环设计，因此能够在保持模块解耦的同时形成产品协同。")

    # Page 5: lesson page.
    doc.add_page_break()
    doc.add_heading("四、学生端证据 1：从动画到代码的同屏学习", level=1)
    add_picture(doc, ASSET_DIR / "lesson-page.png", 6.5,
                "高精度加法课程页：动画、概念标签、理解路径、C++ 示例与课时导航",
                "高精度加法课程页截图", 4)
    doc.add_heading("界面如何服务学习，而不是只追求展示", level=2)
    for text in [
        "动画区把“反向存储、逐位相加、carry 进位、取模与整除”拆成可观察步骤。",
        "概念标签与理解路径帮助学生在观看过程中建立术语和步骤锚点。",
        "C++ 示例与动画共用变量含义，学生可把画面中的变化直接对应到代码。",
        "课时导航、标记完成、练习状态和继续学习入口共同维护学习节奏。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "教学设计意图", "让学生先形成过程模型，再进入语法实现；降低一上来面对整段代码时的认知负荷。")

    # Page 6: judge.
    doc.add_page_break()
    doc.add_heading("五、学生端证据 2：在线编码、判题与复盘", level=1)
    add_picture(doc, ASSET_DIR / "online-judge.png", 6.5,
                "Monaco C++17 编辑器、判题结果与最近提交在同一页面完成",
                "在线 C++ 判题页面截图", 5)
    doc.add_heading("可验证的练习闭环", level=2)
    steps = [
        "学生从课程练习入口进入题目，编辑器自动加载起始代码并保存草稿。",
        "提交后由 FastAPI 调用 Judge Adapter，执行公开与隐藏测试点。",
        "页面展示 Accepted、Wrong Answer、Compile Error 等结果和通过测试点数量。",
        "Accepted 自动计入学习进度；错误提交保留在最近记录中，便于回看和复盘。",
        "分层提示、题解要点和常见错误在卡住时提供逐步帮助，而不是直接暴露完整答案。",
    ]
    for step in steps:
        add_list_item(doc, step, decimal_id)
    add_callout(doc, "工程边界", "当前 local-process 判题器仅用于 MVP 本地演示；正式部署需启用 Docker 沙箱或接入外部 OJ，项目已为此预留 Adapter。", fill="FEF2F2", color=RED)

    # Page 7: dashboard.
    doc.add_page_break()
    doc.add_heading("六、教师端证据：从提交记录到教学诊断", level=1)
    add_picture(doc, ASSET_DIR / "teacher-dashboard.png", 6.5,
                "提交统计、练习覆盖、薄弱题、错误类型和教师建议",
                "教师视角学习数据看板截图", 6)
    doc.add_heading("评委一眼应看到的三个变化", level=2)
    for text in [
        "从“学生做没做”升级为“学生卡在哪里”：薄弱题排行结合低通过率、无 Accepted 和错误提交密度。",
        "从“总体分数”升级为“错误类型”：答案错误、编译错误等分布帮助教师判断是概念、边界还是语法问题。",
        "从“展示数据”升级为“建议行动”：看板根据薄弱题、错误类型和练习覆盖给出下一次课建议。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_body(doc, "当前截图使用可重复生成的 Demo 提交数据，用于稳定展示统计逻辑；后续接入真实账号、班级和日期筛选后，可直接扩展为课堂反馈看板。")

    # Page 8: curriculum.
    doc.add_page_break()
    doc.add_heading("七、课程规模与内容生产证据", level=1)
    add_body(doc, "下表由当前前端课程配置统计，数据对应仓库中的实际课时与练习，不把规划中的空壳能力计入已完成成果。")
    add_curriculum_table(doc)
    p = doc.add_paragraph(style="Caption")
    r = p.add_run("表 1  当前课程覆盖与动画接入情况（203 为课时引用数，去重后为 201 个唯一题目）")
    set_run_font(r, size=9, color=MUTED)
    add_preview_strip(doc)
    add_body(doc, "高精度、排序、递推和递归已形成统一视觉语言；搜索与回溯、贪心、分治、BFS 与动态规划已具备课时与练习结构，其中部分章节配有逐镜头制作脚本，动画将按现有模板持续补齐。")

    # Page 9: competitive advantage + validation.
    doc.add_page_break()
    doc.add_heading("八、参赛优势、可行性与验证", level=1)
    doc.add_heading("相对优势", level=2)
    advantages = [
        "相对录播课程：学生不是被动观看，而是立即进入同知识点的代码与练习。",
        "相对在线题库：系统先建立过程理解，再通过分层提示和题解完成反馈。",
        "相对单个动画 Demo：项目有统一课程结构、201 个唯一题目、判题接口和数据回流。",
        "相对概念型参赛方案：可一键启动，有真实页面、API、提交记录和演示检查脚本。",
    ]
    for item in advantages:
        add_list_item(doc, item, bullet_id)
    doc.add_heading("工程验证", level=2)
    checks = [
        ("前端", "TypeScript 编译与 Vite 生产构建可执行；课程媒体路径与题目 ID 可检查。"),
        ("后端", "FastAPI 健康检查、提交、历史、统计和 Judge 能力查询接口已实现。"),
        ("内容", "Manim 场景提供 480p15 预览与多项 1080p60 成片，预览图直接进入课程。"),
        ("演示", "seed_demo_data.py 可重复生成看板数据，demo_start / stop / check 脚本降低现场风险。"),
    ]
    table = doc.add_table(rows=len(checks) + 1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1850, 7510])
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(["验证维度", "证据"]):
        set_cell_shading(table.cell(0, i), NAVY)
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=WHITE, bold=True)
    for r_i, (label, detail) in enumerate(checks, 1):
        set_cell_shading(table.cell(r_i, 0), GREEN_LIGHT)
        for c_i, text in enumerate((label, detail)):
            p = table.cell(r_i, c_i).paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=GREEN_DARK if c_i == 0 else INK, bold=(c_i == 0))
    doc.add_heading("预期价值", level=2)
    add_body(doc, "对学生，降低算法抽象门槛并缩短从理解到独立编码的距离；对教师，提供可观察的错误与复讲依据；对内容团队，降低跨章节动画与练习生产成本；对平台，提供可独立接入的课程、判题与统计模块。")

    # Page 10: roadmap and demo guide.
    doc.add_page_break()
    doc.add_heading("九、风险控制、路线图与现场演示", level=1)
    doc.add_heading("已识别风险及控制措施", level=2)
    risks = [
        ("判题安全", "本地进程判题仅用于演示；生产环境切换 Docker 沙箱或外部 OJ Adapter。"),
        ("内容完整度", "明确区分 35 个已接入动画课时与尚待成片章节，不把规划能力包装为完成成果。"),
        ("数据真实性", "Demo 数据只用于稳定展示统计逻辑；真实课堂评估将引入账号、班级与时间维度。"),
        ("依赖与部署", "演示前运行检查脚本，后续补充容器化、依赖锁定与持续集成。"),
    ]
    for label, detail in risks:
        add_body(doc, f"{label}：{detail}", bold_lead=f"{label}：")
    doc.add_heading("下一阶段路线图", level=2)
    roadmap = [
        "补齐搜索与回溯、贪心、分治、BFS、动态规划的重点动画成片。",
        "完成 Docker 判题器实测、资源限制、危险系统调用控制与并发队列。",
        "接入真实学生与班级，增加日期、章节和学生维度筛选。",
        "开展课堂试点，以前后测、完成率、平均尝试次数和典型错误变化评估效果。",
        "在行为数据基础上生成个性化复习路径与教师课后复习包。",
    ]
    for item in roadmap:
        add_list_item(doc, item, decimal_id)
    doc.add_heading("建议 3 分钟现场演示路径", level=2)
    demo_steps = [
        "打开高精度加法课时：展示动画与变量映射。",
        "进入“大整数加法”：展示编辑、提交与测试点结果。",
        "返回课程：展示 Accepted 自动计入进度。",
        "打开教师看板：展示薄弱题、错误分布与教学建议。",
    ]
    for item in demo_steps:
        add_list_item(doc, item, decimal_id)
    add_callout(doc, "最终结论", "本项目以可运行工程证明了一条清晰路径：把算法过程可视化，把理解结果转化为代码能力，再把练习数据转化为教学决策。它既具备比赛现场的可展示性，也具备后续真实课堂验证和平台化扩展的基础。")

    doc.core_properties.title = "青少年 C++ 算法可视化教学系统——开题报告附加材料"
    doc.core_properties.subject = "比赛参赛项目佐证材料"
    doc.core_properties.keywords = "C++, 算法可视化, Manim, 在线判题, 教学数据"
    doc.core_properties.comments = "narrative_proposal preset; CJK body: Microsoft YaHei; Chinese headings: SimSun"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
